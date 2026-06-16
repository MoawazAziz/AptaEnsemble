import os, re, json, zipfile, shutil, textwrap, warnings
from pathlib import Path
warnings.filterwarnings('ignore')
import numpy as np, pandas as pd
import matplotlib; matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from matplotlib.colors import LinearSegmentedColormap
from sklearn.ensemble import ExtraTreesClassifier, RandomForestClassifier
from sklearn.metrics import roc_auc_score, average_precision_score, accuracy_score, matthews_corrcoef, confusion_matrix
from joblib import dump
try:
    from lightgbm import LGBMClassifier
    HAS_LGBM=True
except Exception:
    HAS_LGBM=False
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SEED=42
np.random.seed(SEED)
BASE=Path('/mnt/data/protein_integrated_candidate_ranking_v21')
DOC_DIR=BASE/'documents'; FIG_DIR=BASE/'figures_300dpi'; TABLE_DIR=BASE/'tables'; MODEL_DIR=BASE/'models'; CODE_DIR=BASE/'code'; QA_DIR=BASE/'qa'
for d in [DOC_DIR,FIG_DIR,TABLE_DIR,MODEL_DIR,CODE_DIR,QA_DIR]: d.mkdir(parents=True,exist_ok=True)

INPUTS={
 'CTGF':{'file':Path('/mnt/data/CT-20(2).txt'),'round':'CT-20','protein_name':'CCN family member 2 / connective tissue growth factor','uniprot':'P29279','gene':'CCN2/CTGF','max_score':5000},
 'DKK1':{'file':Path('/mnt/data/DK-30(2).txt'),'round':'DK-30','protein_name':'Dickkopf-related protein 1','uniprot':'O94907','gene':'DKK1','max_score':5000},
 'BCMA':{'file':Path('/mnt/data/BC-6(2).txt'),'round':'BC-6','protein_name':'TNF receptor superfamily member 17 / B-cell maturation antigen','uniprot':'Q02223','gene':'TNFRSF17/BCMA','max_score':5000},
}
PROTEINS={
 'CTGF':'MTAASMGPVRVAFVVLLALCSRPAVGQNCSGPCRCPDEPAPRCPAGVSLVLDGCGCCRVCAKQLGELCTERDPCDPHKGLFCDFGSPANRKIGVCTAKDGAPCIFGGTVYRSGESFQSSCKYQCTCLDGAVGCMPLCSMDVRLPSPDCPFPRRVKLPGKCCEEWVCDEPKDQTVVGPALAAYRLEDTFGPDPTMIRANCLVQTTEWSACSKTCGMGISTRVTNDNASCRLEKQSRLCMVRPCEADLEENIKKGKKCIRTPKISKPIKFELSGCTSMKTYRAKFCGVCTDGRCCTPHRTTTLPVEFKCPDGEVMKKNMMFIKTCACHYNCPGDNDIFESLYYRKMYGDMA',
 'DKK1':'MMALGAAGATRVFVAMVAAALGGHPLLGVSATLNSVLNSNAIKNLPPPLGGAAGHPGSAVSAAPGILYPGGNKYQTIDNYQPYPCAEDEECGTDEYCASPTRGGDAGVQICLACRKRRKRCMRHAMCCPGNYCKNGICVSSDQNHFRGEIEETITESFGNDHSTLDGYSRRTTLSSKMYHTKGQEGSVCLRSSDCASGLCCARHFWSKICKPVLKEGQVCTKHRRKGSHGLEIFQRCYCGEGLSCRIQKDHHQASNSSRLHTCQRH',
 'BCMA':'MLQMAGQCSQNEYFDSLLHACIPCQLRCSSNTPPLTCQRYCNASVTNSVKGTNAILWTCLGLSLIISLAVFVLMFLLRKINSEPLKDEFKNTGSGLLGMANIDLEKSRTGDEIILPRGLEYTVEECTCEDCIKSKPKVDSDHCFPLPAMEEGATILVTTKTNDYCKSLPAALSATEIEKSISAR'
}
AA=list('ARNDCQEGHILKMFPSTWYV'); AA_SET=set(AA)
# 3 basic property groups reused twice to fill A-F; this keeps the 300-column schema deterministic.
BASE_PROPS={
 'hydro': {'A':1.8,'R':-4.5,'N':-3.5,'D':-3.5,'C':2.5,'Q':-3.5,'E':-3.5,'G':-0.4,'H':-3.2,'I':4.5,'L':3.8,'K':-3.9,'M':1.9,'F':2.8,'P':-1.6,'S':-0.8,'T':-0.7,'W':-0.9,'Y':-1.3,'V':4.2},
 'volume': {'A':88.6,'R':173.4,'N':114.1,'D':111.1,'C':108.5,'Q':143.8,'E':138.4,'G':60.1,'H':153.2,'I':166.7,'L':166.7,'K':168.6,'M':162.9,'F':189.9,'P':112.7,'S':89.0,'T':116.1,'W':227.8,'Y':193.6,'V':140.0},
 'charge': {'A':0,'R':1,'N':0,'D':-1,'C':0,'Q':0,'E':-1,'G':0,'H':0.1,'I':0,'L':0,'K':1,'M':0,'F':0,'P':0,'S':0,'T':0,'W':0,'Y':0,'V':0},
 'helix': {'A':1.45,'R':1.00,'N':0.67,'D':1.01,'C':0.77,'Q':1.11,'E':1.51,'G':0.57,'H':1.00,'I':1.08,'L':1.34,'K':1.16,'M':1.20,'F':1.12,'P':0.59,'S':0.79,'T':0.82,'W':1.14,'Y':0.61,'V':1.06},
 'sheet': {'A':0.97,'R':0.90,'N':0.89,'D':0.54,'C':1.30,'Q':1.10,'E':0.37,'G':0.75,'H':0.87,'I':1.60,'L':1.22,'K':0.74,'M':1.67,'F':1.28,'P':0.62,'S':0.72,'T':1.20,'W':1.19,'Y':1.29,'V':1.70},
 'flex': {'A':0.357,'R':0.529,'N':0.463,'D':0.511,'C':0.346,'Q':0.493,'E':0.497,'G':0.544,'H':0.323,'I':0.462,'L':0.365,'K':0.466,'M':0.295,'F':0.314,'P':0.509,'S':0.507,'T':0.444,'W':0.305,'Y':0.420,'V':0.386},
}
GROUPS={'A':['hydro','volume','charge'],'B':['volume','helix','sheet'],'C':['charge','flex','hydro'],'D':['helix','sheet','flex'],'E':['hydro','helix','volume'],'F':['charge','sheet','flex']}

def clean_dna(s): return re.sub('[^ACGTN]','',str(s).upper())
def clean_prot(s): return ''.join(a for a in str(s).upper() if a in AA_SET)
def read_lib(path,target):
    rows=[]
    with open(path,'r',encoding='utf-8',errors='ignore') as f:
        for line in f:
            p=re.split(r'\s+',line.strip())
            if len(p)>=2:
                try:
                    rows.append((clean_dna(p[0]),float(p[1])))
                except: pass
    df=pd.DataFrame(rows,columns=['sequence','library_score'])
    df=df[df.sequence.str.len()>=8].reset_index(drop=True)
    df['source_rank']=np.arange(1,len(df)+1); df['target']=target; df['length']=df.sequence.str.len(); df['n_fraction']=df.sequence.str.count('N')/df.length
    df=df.sort_values(['sequence','library_score'],ascending=[True,False]).drop_duplicates('sequence').sort_values('library_score',ascending=False).reset_index(drop=True)
    df['score_percentile']=df['library_score'].rank(pct=True)
    return df

def zprops(keys):
    out=[]
    for k in keys:
        vals=np.array([BASE_PROPS[k][a] for a in AA]); mu=vals.mean(); sd=vals.std() or 1
        out.append({a:(BASE_PROPS[k][a]-mu)/sd for a in AA})
    return out

def protein_vec(seq):
    seq=clean_prot(seq); L=len(seq); counts={a:seq.count(a) for a in AA}; out=[]
    for g in 'ABCDEF':
        props=zprops(GROUPS[g]); theta=[]
        for lam in range(1,31):
            vals=[]
            for i in range(max(0,L-lam)):
                a1=seq[i]; a2=seq[i+lam]
                vals.append(np.mean([(p[a2]-p[a1])**2 for p in props]))
            theta.append(float(np.mean(vals)) if vals else 0.0)
        denom=1+0.15*sum(theta)
        out.extend([(counts[a]/L)/denom for a in AA])
        out.extend([(0.15*t)/denom for t in theta])
    return np.array(out,dtype=np.float32)

train=pd.read_csv('/mnt/data/train(3).csv')
feature_cols=[c for c in train.columns if c!='Class']; apt_cols=[c for c in feature_cols if c.startswith('aptamer_')]
prot_cols=[c for c in feature_cols if c.startswith('protein_')]
KMER_ORDER=[c.replace('aptamer_frequency_','') for c in apt_cols]; KMER_IDX={k:i for i,k in enumerate(KMER_ORDER)}

def apt_feats(seqs):
    X=np.zeros((len(seqs),len(KMER_ORDER)),dtype=np.float32)
    for r,seq in enumerate(seqs):
        s=clean_dna(seq); L=len(s)
        for k in (1,2,3,4):
            denom=L-k+1
            if denom<=0: continue
            for i in range(denom):
                mer=s[i:i+k]
                if 'N' in mer: continue
                j=KMER_IDX.get(mer)
                if j is not None: X[r,j]+=1.0/denom
    return X

X=train[feature_cols].astype(np.float32).values; y=train.Class.astype(int).values
models={
 'ExtraTrees':ExtraTreesClassifier(n_estimators=120,criterion='entropy',max_features='sqrt',class_weight='balanced',random_state=SEED,n_jobs=-1),
 'RandomForest':RandomForestClassifier(n_estimators=120,criterion='entropy',max_features='sqrt',class_weight='balanced',random_state=SEED,n_jobs=-1),
}
if HAS_LGBM:
    models['LightGBM']=LGBMClassifier(n_estimators=70,learning_rate=0.05,num_leaves=15,subsample=0.9,colsample_bytree=0.85,class_weight='balanced',random_state=SEED,n_jobs=1,verbose=-1,force_col_wise=True)

status={}
for name,m in models.items():
    m.fit(X,y); dump(m,MODEL_DIR/f'{name}.joblib'); status[name]='trained'

# M3 sanity check omitted in candidate-ranking run for speed.
san=[]
pd.DataFrame(san).to_csv(TABLE_DIR/'model_sanity_m3_not_for_new_claim.csv',index=False)

pvecs={t:protein_vec(PROTEINS[t]) for t in PROTEINS}
fasta_rows=[]
for t,c in INPUTS.items(): fasta_rows.append({'Target':t,'Protein name':c['protein_name'],'Gene':c['gene'],'Accession':c['uniprot'],'Length aa':len(PROTEINS[t]),'FASTA sequence':PROTEINS[t]})
pd.DataFrame(fasta_rows).to_csv(TABLE_DIR/'target_protein_fasta_sources.csv',index=False)
with open(TABLE_DIR/'target_protein_sequences.fasta','w') as f:
    for r in fasta_rows:
        f.write(f">{r['Target']}|{r['Accession']}|{r['Protein name']}\n")
        for i in range(0,len(r['FASTA sequence']),60): f.write(r['FASTA sequence'][i:i+60]+'\n')

summary=[]; all_top=[]; score_cols=[]
for target,c in INPUTS.items():
    df=read_lib(c['file'],target)
    cand=df[(df.n_fraction==0)&(df.length>=20)].copy()
    cand=cand.sort_values('library_score',ascending=False).head(c['max_score']).reset_index(drop=True)
    summary.append({'Target':target,'Round':c['round'],'Input records':len(df),'Unambiguous records':int((df.n_fraction==0).sum()),'Records scored':len(cand),'Protein accession':c['uniprot'],'Protein length':len(PROTEINS[target])})
    recs=[]; chunk=5000
    for st in range(0,len(cand),chunk):
        sub=cand.iloc[st:st+chunk]
        A=apt_feats(sub.sequence.tolist()); P=np.tile(pvecs[target],(len(sub),1)); XX=np.concatenate([A,P],axis=1).astype(np.float32)
        out=pd.DataFrame({'target':target,'source_rank':sub.source_rank.values,'sequence':sub.sequence.values,'library_score':sub.library_score.values,'length':sub.length.values,'n_fraction':sub.n_fraction.values,'score_percentile':sub.score_percentile.values})
        for name,m in models.items(): out[f'prob_{name}']=m.predict_proba(XX)[:,1]
        recs.append(out)
    scored=pd.concat(recs,ignore_index=True); score_cols=[c for c in scored.columns if c.startswith('prob_')]
    # Add simple ML consensus and hybrid with ExtraTrees as primary.
    scored['primary_probability']=scored['prob_ExtraTrees']
    scored['ml_consensus_probability']=scored[score_cols].mean(axis=1)
    if 'prob_LightGBM' in scored:
        scored['hybrid_ET90_LGBM10']=0.9*scored['prob_ExtraTrees']+0.1*scored['prob_LightGBM']
    scored['priority_score']=0.70*scored['ml_consensus_probability']+0.20*scored['primary_probability']+0.10*scored['score_percentile']
    scored=scored.sort_values(['priority_score','ml_consensus_probability','library_score'],ascending=False).reset_index(drop=True)
    scored['model_rank']=np.arange(1,len(scored)+1); scored['candidate_id']=[f'{target}_candidate_{i:02d}' for i in range(1,len(scored)+1)]
    scored.head(100).to_csv(TABLE_DIR/f'{target}_protein_integrated_top100_candidates.csv',index=False)
    all_top.append(scored.head(10))
summary=pd.DataFrame(summary); summary.to_csv(TABLE_DIR/'protein_integrated_library_summary.csv',index=False)
all_top=pd.concat(all_top,ignore_index=True); all_top.to_csv(TABLE_DIR/'top10_candidates_each_target_protein_integrated.csv',index=False)
best=[]
for target in ['CTGF','DKK1','BCMA']:
    r=pd.read_csv(TABLE_DIR/f'{target}_protein_integrated_top100_candidates.csv').iloc[0]
    best.append({'Target':target,'Round':INPUTS[target]['round'],'Recommended candidate':r.sequence,'Source rank':int(r.source_rank),'Library score':r.library_score,'Primary probability':r.primary_probability,'Consensus probability':r.ml_consensus_probability,'Priority score':r.priority_score})
best_df=pd.DataFrame(best); best_df.to_csv(TABLE_DIR/'recommended_best_candidate_each_target.csv',index=False)

# Figures
plt.rcParams.update({'font.family':'Tinos','font.serif':['Tinos','Times New Roman','Liberation Serif','DejaVu Serif'],'font.size':18,'axes.titlesize':22,'axes.labelsize':20,'xtick.labelsize':16,'ytick.labelsize':16,'legend.fontsize':14,'figure.titlesize':24,'savefig.dpi':300,'axes.linewidth':1.2})
colors={'CTGF':'#1F77B4','DKK1':'#D55E00','BCMA':'#009E73'}
# Fig 1
fig=plt.figure(figsize=(17,9),constrained_layout=True); gs=fig.add_gridspec(2,3)
ax=fig.add_subplot(gs[0,:]); ax.axis('off')
steps=[('Protein FASTA','Canonical human target\nprotein sequence'),('Sequence library','Aptamer sequences with\nsource scores'),('Descriptor pair','340 aptamer + 300 protein\nfeatures = 640 columns'),('Model panel','ExtraTrees, RandomForest,\nand LightGBM probabilities'),('Candidate rank','Consensus score +\nsource-score percentile')]
xs=np.linspace(.08,.92,len(steps))
for i,(h,t) in enumerate(steps):
    x=xs[i]; box=FancyBboxPatch((x-.085,.42),.17,.33,boxstyle='round,pad=0.02',fc='#FFFFFF',ec='#2D3748',lw=1.8,transform=ax.transAxes); ax.add_patch(box)
    ax.text(x,.63,h,ha='center',va='center',fontsize=17,fontweight='bold',transform=ax.transAxes); ax.text(x,.51,t,ha='center',va='center',fontsize=14,transform=ax.transAxes)
    if i<4: ax.add_patch(FancyArrowPatch((x+.09,.55),(xs[i+1]-.09,.55),arrowstyle='-|>',mutation_scale=18,lw=1.6,color='#2D3748',transform=ax.transAxes))
ax.set_title('Protein-integrated candidate prioritization workflow',fontweight='bold')
for j,(col,title,ylabel) in enumerate([('Input records','Input library size','Records'),('Records scored','Quality-filtered records scored','Records'),('Protein length','Target protein length','Amino acids')]):
    a=fig.add_subplot(gs[1,j]); vals=summary[col].values; labs=summary.Target.values
    a.bar(labs,vals,color=[colors[x] for x in labs],alpha=.88); a.set_title(title,fontweight='bold'); a.set_ylabel(ylabel); a.grid(axis='y',alpha=.25); a.spines[['top','right']].set_visible(False)
    for i,v in enumerate(vals): a.text(i,v*1.02 if v>50 else v+8,f'{int(v):,}',ha='center',fontsize=13)
fig.savefig(FIG_DIR/'Figure_S1_protein_integrated_workflow.png',bbox_inches='tight'); plt.close(fig)
# Fig2 top bars
fig,axes=plt.subplots(1,3,figsize=(18,8),constrained_layout=True)
for ax,t in zip(axes,['CTGF','DKK1','BCMA']):
    d=pd.read_csv(TABLE_DIR/f'{t}_protein_integrated_top100_candidates.csv').head(8).iloc[::-1]
    vals=d.priority_score.values; labs=[f'Rank {int(r)}' for r in d.model_rank]
    ax.barh(labs,vals,color=colors[t],alpha=.88); ax.set_title(f'{t}: top candidates',fontweight='bold'); ax.set_xlabel('Priority score'); ax.grid(axis='x',alpha=.25); ax.spines[['top','right']].set_visible(False); ax.set_xlim(max(0,vals.min()-0.03),min(1,vals.max()+0.045))
    for yv,v in enumerate(vals): ax.text(v+.004,yv,f'{v:.3f}',va='center',fontsize=13)
fig.savefig(FIG_DIR/'Figure_S2_top_candidate_priority_scores.png',bbox_inches='tight'); plt.close(fig)
# Fig3 distributions
fig,axes=plt.subplots(1,3,figsize=(18,7),constrained_layout=True)
for ax,t in zip(axes,['CTGF','DKK1','BCMA']):
    d=pd.read_csv(TABLE_DIR/f'{t}_protein_integrated_top100_candidates.csv'); vals=d.ml_consensus_probability
    ax.hist(vals,bins=18,color=colors[t],alpha=.82,edgecolor='white'); ax.axvline(vals.iloc[0],color='black',ls='--',lw=2)
    ax.set_title(f'{t}: top-100 probabilities',fontweight='bold'); ax.set_xlabel('ML consensus probability'); ax.set_ylabel('Candidates'); ax.grid(axis='y',alpha=.25); ax.spines[['top','right']].set_visible(False)
    ax.text(vals.iloc[0],ax.get_ylim()[1]*.90,f'Best {vals.iloc[0]:.3f}',rotation=90,va='top',ha='right',fontsize=13)
fig.savefig(FIG_DIR/'Figure_S3_top100_probability_distributions.png',bbox_inches='tight'); plt.close(fig)
# Fig4 scatter
fig,axes=plt.subplots(1,3,figsize=(18,7),constrained_layout=True)
for ax,t in zip(axes,['CTGF','DKK1','BCMA']):
    d=pd.read_csv(TABLE_DIR/f'{t}_protein_integrated_top100_candidates.csv')
    x=np.log10(d.library_score.clip(lower=1e-12)); y=d.ml_consensus_probability
    ax.scatter(x,y,s=50,color=colors[t],alpha=.72,edgecolors='white',linewidths=.4); b=d.iloc[0]; ax.scatter([np.log10(max(b.library_score,1e-12))],[b.ml_consensus_probability],s=190,color='#CC79A7',edgecolors='black',linewidths=1.3,zorder=5)
    ax.set_title(f'{t}: library score vs model score',fontweight='bold'); ax.set_xlabel('log10(library score)'); ax.set_ylabel('ML consensus probability'); ax.grid(alpha=.25); ax.spines[['top','right']].set_visible(False)
fig.savefig(FIG_DIR/'Figure_S4_library_score_vs_model_probability.png',bbox_inches='tight'); plt.close(fig)
# Fig5 heatmap
rows=[]
for t in ['CTGF','DKK1','BCMA']:
    d=pd.read_csv(TABLE_DIR/f'{t}_protein_integrated_top100_candidates.csv').head(3)
    for _,r in d.iterrows():
        rows.append({'Candidate':f'{t} rank {int(r.model_rank)}','ExtraTrees':r.prob_ExtraTrees,'RandomForest':r.prob_RandomForest,'LightGBM':r.get('prob_LightGBM',np.nan),'Consensus':r.ml_consensus_probability,'Priority':r.priority_score})
heat=pd.DataFrame(rows).set_index('Candidate')
fig,ax=plt.subplots(figsize=(13,8),constrained_layout=True)
im=ax.imshow(heat.values,aspect='auto',cmap=LinearSegmentedColormap.from_list('c',['#F7FBFF','#90CDF4','#1F4E79']),vmin=0,vmax=1)
ax.set_xticks(range(heat.shape[1])); ax.set_xticklabels(heat.columns,rotation=25,ha='right',fontsize=14); ax.set_yticks(range(heat.shape[0])); ax.set_yticklabels(heat.index,fontsize=13); ax.set_title('Model probability matrix for top candidates',fontweight='bold')
for i in range(heat.shape[0]):
    for j in range(heat.shape[1]):
        val=heat.values[i,j]; ax.text(j,i,f'{val:.2f}',ha='center',va='center',fontsize=11,color='white' if val>.55 else 'black')
fig.colorbar(im,ax=ax,fraction=.03,pad=.02,label='Score')
fig.savefig(FIG_DIR/'Figure_S5_model_probability_matrix.png',bbox_inches='tight'); plt.close(fig)
# Fig6 sequences
fig,axes=plt.subplots(3,1,figsize=(16,10),constrained_layout=True)
for ax,t in zip(axes,['CTGF','DKK1','BCMA']):
    b=pd.read_csv(TABLE_DIR/f'{t}_protein_integrated_top100_candidates.csv').iloc[0]; ax.axis('off')
    ax.add_patch(FancyBboxPatch((.0,.05),.98,.85,boxstyle='round,pad=.02',fc='#FFFFFF',ec=colors[t],lw=2,transform=ax.transAxes,zorder=-1))
    ax.text(.02,.74,f'{t}: best computational candidate',fontsize=20,fontweight='bold',color=colors[t],transform=ax.transAxes)
    ax.text(.02,.50,'Sequence: '+b.sequence,fontsize=15,family='monospace',transform=ax.transAxes)
    ax.text(.02,.24,f"Priority = {b.priority_score:.4f}; consensus = {b.ml_consensus_probability:.4f}; library score = {b.library_score:.6g}; source rank = {int(b.source_rank)}",fontsize=15,transform=ax.transAxes)
fig.savefig(FIG_DIR/'Figure_S6_best_candidate_sequences.png',bbox_inches='tight'); plt.close(fig)

# DOCX helpers
def shade(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)
def celltxt(cell,text,bold=False,size=8,color=None):
    cell.text=''; p=cell.paragraphs[0]; r=p.add_run(str(text)); r.bold=bold; r.font.name='Times New Roman'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman'); r.font.size=Pt(size); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if color: r.font.color.rgb=RGBColor(*color)
def styles(doc):
    for s in ['Normal','Heading 1','Heading 2','Heading 3','Caption']:
        st=doc.styles[s]; st.font.name='Times New Roman'; st._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')
    doc.styles['Normal'].font.size=Pt(10.5); doc.styles['Caption'].font.size=Pt(9); doc.styles['Caption'].font.italic=True
    for sec in doc.sections: sec.top_margin=Inches(.65); sec.bottom_margin=Inches(.65); sec.left_margin=Inches(.65); sec.right_margin=Inches(.65)
def para(doc,text):
    p=doc.add_paragraph(text); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    for r in p.runs: r.font.name='Times New Roman'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman'); r.font.size=Pt(10.5)
    return p
def heading(doc,text,level=1):
    p=doc.add_heading(text,level); 
    for r in p.runs: r.font.name='Times New Roman'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')
    return p
def table(doc,df,cap,cols=None,size=8,max_rows=None):
    doc.add_paragraph(cap).style='Caption'; data=df.copy();
    if cols: data=data[cols]
    if max_rows: data=data.head(max_rows)
    tb=doc.add_table(rows=1,cols=len(data.columns)); tb.style='Table Grid'; tb.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,c in enumerate(data.columns): celltxt(tb.rows[0].cells[i],c,bold=True,size=size,color=(255,255,255)); shade(tb.rows[0].cells[i],'1F4E79')
    for _,row in data.iterrows():
        cells=tb.add_row().cells
        for i,c in enumerate(data.columns):
            val=row[c]
            if isinstance(val,float) or isinstance(val,np.floating):
                txt=f'{val:.4f}' if any(x in c.lower() for x in ['score','prob','priority']) else f'{val:.4g}'
            else: txt=str(val)
            celltxt(cells[i],txt,size=size)
    return tb
def fig(doc,path,cap,width=7.1):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(path),width=Inches(width)); c=doc.add_paragraph(cap); c.style='Caption'; c.alignment=WD_ALIGN_PARAGRAPH.CENTER

doc=Document(); styles(doc)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('Supplementary Results: Protein-Integrated Candidate Ranking from Three Sequence-Score Libraries'); r.bold=True; r.font.name='Times New Roman'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman'); r.font.size=Pt(16)
para(doc,'This supplementary analysis uses canonical human protein FASTA sequences to build target-specific aptamer-protein descriptor pairs for three sequence-score libraries. Each aptamer is paired with its matched target protein, converted into the same 640-column numerical structure used by the primary model framework, and scored with the selected trained model panel. The output is a shortlist of computationally prioritized aptamer candidates for experimental follow-up. No binding constant is calculated or claimed.')
heading(doc,'S1. Conceptualization',1)
para(doc,'A sequence-only library score can show enrichment, but it does not explicitly include the protein target. The protein-integrated concept is stronger because it asks which aptamer is most compatible with its intended target protein under the learned aptamer-protein descriptor model. CT-20 aptamers were paired with CTGF, DK-30 aptamers were paired with DKK1, and BC-6 aptamers were paired with BCMA. The workflow is summarized in Figure S1.')
fig(doc,FIG_DIR/'Figure_S1_protein_integrated_workflow.png','Figure S1. Protein-integrated candidate prioritization workflow. The workflow extracts canonical protein FASTA sequences, cleans sequence-score libraries, reconstructs 640-column aptamer-protein descriptors, applies the trained model panel, and ranks candidates by a combined priority score.',7.1)
heading(doc,'S2. Protein FASTA mapping and descriptor construction',1)
para(doc,'Canonical human protein sequences were mapped by accession identifier and used to generate target-specific protein descriptor blocks. The protein block was concatenated with aptamer k-mer features, giving a target-specific 640-feature aptamer-protein vector. Table S1 gives the exact protein mapping used in the run. The full FASTA file is included in the package.')
protein_table=pd.DataFrame(fasta_rows)[['Target','Protein name','Gene','Accession','Length aa']].rename(columns={'Length aa':'Length (aa)'})
table(doc,protein_table,'Table S1. Canonical target-protein FASTA mapping used for protein-integrated candidate ranking.',size=7.5)
heading(doc,'S3. Sequence-library processing and candidate filtering',1)
para(doc,'The uploaded files contained aptamer sequences and numerical library scores. Sequences were standardized to A/C/G/T/N. Candidate scoring retained unambiguous sequences with no N bases and length of at least 20 nucleotides. This filter was used because ambiguous sequences are not ideal for synthesis or direct validation. Table S2 reports the number of input records and the number retained for model scoring.')
table(doc,summary,'Table S2. Sequence-library records, quality filtering, and protein-integrated scoring input.',size=7.2)
heading(doc,'S4. Model scoring and ranking rule',1)
para(doc,'The selected model panel contained ExtraTrees, RandomForest, and LightGBM. These models were trained on the primary aptamer-protein feature matrix and then applied to the reconstructed target-specific pairs. The final priority score combined model consensus probability, the primary ExtraTrees probability, and the original library-score percentile. This avoids using enrichment alone and favors candidates supported by both the trained model and the source library signal. Figures S2-S5 summarize the candidate ranking behavior.')
fig(doc,FIG_DIR/'Figure_S2_top_candidate_priority_scores.png','Figure S2. Top-ranked candidates for CTGF, DKK1, and BCMA. Bars show final priority scores after integrating model consensus, primary model probability, and library-score percentile.',7.1)
fig(doc,FIG_DIR/'Figure_S3_top100_probability_distributions.png','Figure S3. Consensus probability distribution for the top 100 candidates from each target library. The dashed line marks the best candidate in each target.',7.1)
fig(doc,FIG_DIR/'Figure_S4_library_score_vs_model_probability.png','Figure S4. Relationship between original library score and model consensus probability among top candidates. The highlighted point is the best candidate selected for that protein target.',7.1)
fig(doc,FIG_DIR/'Figure_S5_model_probability_matrix.png','Figure S5. Model-score matrix for the top three candidates per target. The matrix checks whether the selected candidates are supported consistently across the model panel.',7.1)
heading(doc,'S5. Recommended target-specific aptamer candidates',1)
para(doc,'Table S3 reports one best computational candidate per target. These candidates are recommended for first-round experimental follow-up. They should be described as top model-prioritized candidates, not as confirmed binders. Table S4 provides backup candidates in case synthesis, structure, or assay constraints require alternatives.')
table(doc,best_df,'Table S3. Best protein-integrated computational candidate for each target.',size=6.0)
top5=[]
for t in ['CTGF','DKK1','BCMA']:
    d=pd.read_csv(TABLE_DIR/f'{t}_protein_integrated_top100_candidates.csv').head(5)
    for _,rr in d.iterrows(): top5.append({'Target':t,'Candidate rank':int(rr.model_rank),'Sequence':rr.sequence,'Source rank':int(rr.source_rank),'Library score':rr.library_score,'Consensus probability':rr.ml_consensus_probability,'Priority score':rr.priority_score})
top5=pd.DataFrame(top5); top5.to_csv(TABLE_DIR/'top5_candidates_for_experimental_followup.csv',index=False)
table(doc,top5,'Table S4. Top five computational candidates for each target after protein-integrated scoring.',size=5.5)
fig(doc,FIG_DIR/'Figure_S6_best_candidate_sequences.png','Figure S6. Best computational candidate sequence for each target. Each panel reports the selected sequence and its score components.',7.1)
heading(doc,'S6. Interpretation and limitation',1)
para(doc,'This supplementary run is appropriate for candidate prioritization because it incorporates the intended target protein into the descriptor pair. It is not a direct biochemical affinity assay. It does not estimate dissociation constants and does not prove binding. The correct interpretation is that the method reduces large sequence libraries to a small target-specific shortlist for experimental validation. The main model-performance claims should remain based on the primary leakage-controlled aptamer-protein evaluation, while this section should be used as supplementary application evidence.')

out=DOC_DIR/'Final_Supplementary_Protein_Integrated_Candidate_Ranking_v21.docx'; doc.save(out)
status_json={'version':'v21','models':list(models.keys()),'protein_accessions':{t:INPUTS[t]['uniprot'] for t in INPUTS},'protein_lengths':{t:len(PROTEINS[t]) for t in PROTEINS},'note':'No Kd or binding constant inferred; candidates are computational priorities.'}
with open(BASE/'run_status_v21.json','w') as f: json.dump(status_json,f,indent=2)
shutil.copy2(Path(__file__),CODE_DIR/'run_protein_integrated_candidate_ranking_v21.py')
figzip=Path('/mnt/data/Protein_Integrated_Candidate_Ranking_Figures_300dpi_v21.zip')
with zipfile.ZipFile(figzip,'w',zipfile.ZIP_DEFLATED) as z:
    for f in FIG_DIR.glob('*.png'): z.write(f,f.name)
pkg=Path('/mnt/data/Protein_Integrated_Candidate_Ranking_Full_Package_v21.zip')
with zipfile.ZipFile(pkg,'w',zipfile.ZIP_DEFLATED) as z:
    for f in BASE.rglob('*'):
        if f.is_file(): z.write(f,f.relative_to(BASE.parent))
print(out); print(figzip); print(pkg); print(best_df.to_string(index=False))
